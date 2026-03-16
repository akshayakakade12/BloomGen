import streamlit as st
from langchain_groq import ChatGroq
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_text_splitters import RecursiveCharacterTextSplitter
import os
import math
import docx
import PyPDF2
import time
from docx import Document
import datetime
from dotenv import load_dotenv

# DOCX formatting helpers
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import pandas as pd

load_dotenv()

# =========================
# LOGIN SYSTEM
# =========================
users = {
    "admin": {"password": "admin123", "role": "Admin"},
    "educator": {"password": "edu123", "role": "Educator"},
    "student": {"password": "stu123", "role": "Student"}
}

if "logged_in" not in st.session_state:
    st.session_state.logged_in = False

if "role" not in st.session_state:
    st.session_state.role = None

if "mode" not in st.session_state:
    st.session_state.mode = None

if "preview_rows" not in st.session_state:
    st.session_state.preview_rows = None

if "generated_docx_path" not in st.session_state:
    st.session_state.generated_docx_path = None

# =========================
# LOGIN PAGE
# =========================
if not st.session_state.logged_in:

    st.title("🌸 BloomGen - Login")

    username = st.text_input("Username")
    password = st.text_input("Password", type="password")

    if st.button("Login"):
        if username in users and users[username]["password"] == password:
            st.session_state.logged_in = True
            st.session_state.role = users[username]["role"]
            st.success(f"Logged in as {st.session_state.role}")
            st.rerun()
        else:
            st.error("Invalid credentials")

# =========================
# DASHBOARD
# =========================
else:
    st.title("🌸 BloomGen - AI Academic Generator")
    st.sidebar.success(f"Logged in as {st.session_state.role}")

    if st.sidebar.button("Logout"):
        st.session_state.logged_in = False
        st.session_state.role = None
        st.session_state.mode = None
        st.session_state.preview_rows = None
        st.session_state.generated_docx_path = None
        st.rerun()

    if st.session_state.mode is not None:
        if st.sidebar.button("⬅ Back to Home"):
            st.session_state.mode = None
            st.session_state.preview_rows = None
            st.session_state.generated_docx_path = None
            st.rerun()

    # =========================
    # HOME PAGE
    # =========================
    if st.session_state.mode is None:
        st.markdown(
            """
            <h3 style='text-align:center;'>Hiii welcome back, what you want to do today?</h3>
            """,
            unsafe_allow_html=True
        )

        st.write("")
        col1, col2, col3 = st.columns([1, 1.5, 1])

        with col2:
            if st.button("Assignment generation", use_container_width=True):
                st.session_state.mode = "assignment"
                st.rerun()

            st.write("")
            st.write("")

            if st.button("Question paper generation", use_container_width=True):
                st.session_state.mode = "question_paper"
                st.rerun()

    # =========================
    # COMMON SETUP
    # =========================
    else:
        llm = ChatGroq(
            groq_api_key=os.getenv("GROQ_API_KEY"),
            model_name="openai/gpt-oss-120b",
            temperature=0.4,
            max_tokens=1200
        )

        # =========================
        # Safe Retry Wrapper
        # =========================
        def safe_llm_invoke(chain, payload, retries=3, delay=2):
            for attempt in range(retries):
                try:
                    return chain.invoke(payload)
                except Exception as e:
                    error_text = str(e).lower()
                    if (
                        "503" in error_text
                        or "over capacity" in error_text
                        or "internalservererror" in error_text
                        or "service unavailable" in error_text
                    ):
                        if attempt < retries - 1:
                            wait_time = delay * (2 ** attempt)
                            time.sleep(wait_time)
                            continue
                    raise

        # =========================
        # File Text Extraction
        # =========================
        def extract_text(uploaded_file):
            text = ""
            if uploaded_file is None:
                return text

            filename = uploaded_file.name.lower()

            if filename.endswith(".pdf"):
                reader = PyPDF2.PdfReader(uploaded_file)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"

            elif filename.endswith(".docx"):
                doc = docx.Document(uploaded_file)
                for para in doc.paragraphs:
                    text += para.text + "\n"

            return text

        # =========================
        # Syllabus Chunking Helpers
        # =========================
        def split_syllabus(text: str, chunk_size: int = 2400, chunk_overlap: int = 200):
            splitter = RecursiveCharacterTextSplitter(
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap
            )
            return splitter.split_text(text or "")

        def safe_join(parts, sep="\n\n"):
            return sep.join([p for p in parts if p and p.strip()])

        # =========================
        # Summary Chain
        # =========================
        summary_prompt = ChatPromptTemplate.from_template("""
You are helping create university exam/assignment questions.

Summarize the syllabus into compact bullet points (max 350 words).
Keep only exam-relevant units/topics/subtopics/keywords.
No extra explanation.

SYLLABUS:
{syllabus}
""")
        summary_chain = summary_prompt | llm | StrOutputParser()

        # =========================
        # Bloom Count Helper
        # =========================
        def compute_bloom_counts(total_questions: int, pct_u: int, pct_a: int, pct_ae: int):
            u = round(total_questions * pct_u / 100)
            a = round(total_questions * pct_a / 100)
            ae = total_questions - u - a
            return {
                "Understand": max(0, u),
                "Apply": max(0, a),
                "Analyze/Evaluate": max(0, ae)
            }

        # =========================
        # Assignment Question Generator
        # =========================
        def generate_questions(subject, syllabus, count, pct_u, pct_a, pct_ae):
            chunks = split_syllabus(syllabus, chunk_size=2400, chunk_overlap=200)

            summaries = []
            for ch in chunks[:6]:
                summaries.append(safe_llm_invoke(summary_chain, {"syllabus": ch}))
            syllabus_summary = safe_join(summaries)

            bloom_counts = compute_bloom_counts(count, pct_u, pct_a, pct_ae)

            bloom_prompt = ChatPromptTemplate.from_template("""
You are an academic question paper setter.

Generate exactly {count} university-level descriptive questions.

Subject: {subject}
Bloom Bucket: {bloom_bucket}

STRICT OUTPUT RULES:
- Output ONLY questions.
- One question per line.
- No numbering, no bullets, no headings, no blank lines.
- Keep questions exam-oriented and concise.
- Use action verbs that match the Bloom bucket:
  - Understand: Explain, Describe, Illustrate, Summarize
  - Apply: Solve, Demonstrate, Implement, Apply
  - Analyze/Evaluate: Analyze, Compare, Differentiate, Evaluate, Justify

Syllabus Summary (use ONLY this):
{syllabus}
""")
            bloom_chain = bloom_prompt | llm | StrOutputParser()

            batch_size = 6
            final_pairs = []

            for bloom_bucket, bucket_count in bloom_counts.items():
                if bucket_count <= 0:
                    continue

                rounds = math.ceil(bucket_count / batch_size)

                for i in range(rounds):
                    this_batch = batch_size if i < rounds - 1 else (bucket_count - batch_size * i)

                    out = safe_llm_invoke(
                        bloom_chain,
                        {
                            "subject": subject,
                            "syllabus": syllabus_summary,
                            "count": this_batch,
                            "bloom_bucket": bloom_bucket
                        }
                    )

                    for line in out.split("\n"):
                        q = line.strip().lstrip("-•").strip()
                        if q:
                            final_pairs.append((q, bloom_bucket))

            return final_pairs[:count]

        # =========================
        # Question Paper Generators
        # =========================
        def generate_section_questions(subject, unit_text, count, marks, style, bloom_hint):
            chunks = split_syllabus(unit_text, chunk_size=2200, chunk_overlap=150)

            summaries = []
            for ch in chunks[:2]:
                summaries.append(safe_llm_invoke(summary_chain, {"syllabus": ch}))

            unit_summary = safe_join(summaries)

            prompt = ChatPromptTemplate.from_template("""
You are a university exam question setter.

Generate exactly {count} questions.

Subject: {subject}
Marks per Question: {marks}
Question Style: {style}
Bloom Level: {bloom_hint}

RULES:
- Output only questions
- One question per line
- No numbering
- No bullets
- No headings
- Questions must be suitable for {marks} marks
- Keep the style aligned to {style}

Unit Summary:
{unit_summary}
""")

            chain = prompt | llm | StrOutputParser()

            result = safe_llm_invoke(
                chain,
                {
                    "subject": subject,
                    "count": count,
                    "marks": marks,
                    "style": style,
                    "bloom_hint": bloom_hint,
                    "unit_summary": unit_summary
                }
            )

            questions = []
            for line in result.split("\n"):
                q = line.strip().lstrip("-•").strip()
                if q:
                    questions.append(q)

            return questions[:count]

        # =========================
        # CO / PO Helpers
        # =========================
        def assign_co(index: int, total_cos: int) -> str:
            return f"CO{(index % total_cos) + 1}" if total_cos > 0 else ""

        def assign_po(index: int, total_pos: int) -> str:
            return f"PO{(index % total_pos) + 1}" if total_pos > 0 else ""

        def marks_for_bloom(bloom_label: str, m_u: int, m_a: int, m_ae: int) -> str:
            if bloom_label == "Understand":
                return str(m_u)
            if bloom_label == "Apply":
                return str(m_a)
            return str(m_ae)

        # =========================
        # DOCX Formatting Helpers
        # =========================
        DOC_FONT_NAME = "Calibri"
        DOC_FONT_SIZE_PT = 11
        ROW_HEIGHT_PT = 24

        def set_cell_text(cell, text: str, bold: bool = False):
            cell.text = ""
            p = cell.paragraphs[0]

            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p.paragraph_format.line_spacing = 1.0

            run = p.add_run(text)
            run.bold = bold
            run.font.name = DOC_FONT_NAME
            run.font.size = Pt(DOC_FONT_SIZE_PT)

            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.append(rFonts)
            rFonts.set(qn("w:ascii"), DOC_FONT_NAME)
            rFonts.set(qn("w:hAnsi"), DOC_FONT_NAME)

            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        def set_row_height(row, height_pt: int):
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            trHeight = OxmlElement("w:trHeight")
            trHeight.set(qn("w:val"), str(int(height_pt * 20)))
            trHeight.set(qn("w:hRule"), "exact")
            trPr.append(trHeight)

        # =========================
        # Assignment DOCX Generator
        # =========================
        def generate_university_docx(
            data_dict,
            questions_list,
            bloom_labels,
            total_cos,
            total_pos,
            m_u,
            m_a,
            m_ae
        ):
            BASE_DIR = os.path.dirname(os.path.abspath(__file__))
            template_path = os.path.join(BASE_DIR, "templates", "assignment_template.docx")

            doc = Document(template_path)

            for paragraph in doc.paragraphs:
                for key, value in data_dict.items():
                    if key in paragraph.text:
                        paragraph.text = paragraph.text.replace(key, str(value))

            question_table = None
            for table in doc.tables:
                if table.rows and "Question No." in table.rows[0].cells[0].text:
                    question_table = table
                    break

            if question_table:
                q_no = 1
                for idx, q in enumerate(questions_list):
                    q = q.strip()
                    if not q:
                        continue

                    bloom_label = bloom_labels[idx] if idx < len(bloom_labels) else "Understand"

                    new_row = question_table.add_row()
                    set_row_height(new_row, ROW_HEIGHT_PT)
                    cells = new_row.cells

                    set_cell_text(cells[0], f"Q{q_no}")
                    set_cell_text(cells[1], q)
                    set_cell_text(cells[2], assign_co(idx, total_cos))
                    set_cell_text(cells[3], assign_po(idx, total_pos))
                    set_cell_text(cells[4], bloom_label)
                    set_cell_text(cells[5], marks_for_bloom(bloom_label, m_u, m_a, m_ae))

                    q_no += 1

            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M")
            safe_subject = "".join(
                [c for c in (data_dict.get("{{SUBJECT}}", "") or "Subject") if c.isalnum() or c in (" ", "_", "-")]
            ).strip()
            safe_subject = safe_subject.replace(" ", "_")[:40] or "Subject"

            output_path = os.path.join(BASE_DIR, f"University_Assignment_{safe_subject}_{timestamp}.docx")
            doc.save(output_path)
            return output_path

        # =========================
        # Question Paper DOCX Generator
        # =========================
        def generate_question_paper_docx(data_dict, all_rows):
            BASE_DIR = os.path.dirname(os.path.abspath(__file__))
            template_path = os.path.join(BASE_DIR, "templates", "question_paper_template.docx")

            doc = Document(template_path)

            replacements = {
                "{{DEPARTMENT}}": data_dict.get("department", ""),
                "{{SEMESTER}}": data_dict.get("semester", ""),
                "{{ACADEMIC_YEAR}}": data_dict.get("academic_year", ""),
                "{{MAX_MARKS}}": data_dict.get("total_marks", ""),
                "{{COURSE_NAME}}": data_dict.get("course_name", ""),
                "{{COURSE_CODE}}": data_dict.get("course_code", ""),
                "{{SUBJECT_TEACHER}}": data_dict.get("subject_teacher", ""),
                "{{DURATION}}": data_dict.get("duration", ""),
                "{{DATE}}": str(data_dict.get("date", "")),
                "{{EXAMINATION_DATE}}": str(data_dict.get("exam_date", "")),
                "{{SUBJECT_INCHARGE}}": data_dict.get("subject_incharge", ""),
                "{{ACADEMIC_COORDINATOR}}": data_dict.get("academic_coordinator", ""),
                "{{HOD_NAME}}": data_dict.get("hod_name", "")
            }

            # Replace header placeholders in paragraphs
            for paragraph in doc.paragraphs:
                for key, value in replacements.items():
                    if key in paragraph.text:
                        paragraph.text = paragraph.text.replace(key, str(value))

            # Replace placeholders inside tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for key, value in replacements.items():
                                if key in paragraph.text:
                                    paragraph.text = paragraph.text.replace(key, str(value))

            # Split rows by section
            sections = {
                "A": [],
                "B": [],
                "C": [],
                "D": [],
                "E": []
            }

            for row in all_rows:
                sections[row["Section"]].append(row)

            # Assumes:
            # table 0 = logo/header
            # table 1 = meta info
            # table 2 = section A table
            # table 3 = section B table
            # table 4 = section C table
            # table 5 = section D table
            # table 6 = section E table
            section_tables = doc.tables[2:]
            section_order = ["A", "B", "C", "D", "E"]

            for i, sec in enumerate(section_order):
                if i >= len(section_tables):
                    break

                table = section_tables[i]

                for row_data in sections[sec]:
                    new_row = table.add_row()
                    set_row_height(new_row, ROW_HEIGHT_PT)
                    cells = new_row.cells

                    set_cell_text(cells[0], row_data["Question No."])
                    set_cell_text(cells[1], row_data["Question Statement"])
                    set_cell_text(cells[2], row_data["CO"])
                    set_cell_text(cells[3], row_data["PO"])
                    set_cell_text(cells[4], row_data["Bloom’s Level"])
                    set_cell_text(cells[5], row_data["Marks"])

            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M")
            output_path = os.path.join(BASE_DIR, f"Question_Paper_{timestamp}.docx")
            doc.save(output_path)
            return output_path

        # =========================
        # ASSIGNMENT MODE
        # =========================
        if st.session_state.mode == "assignment":
            st.header("📘 Assignment Generation")

            st.sidebar.header("Assignment Details")

            subject = st.sidebar.text_input("Subject")
            department = st.sidebar.text_input("Department", "CSE")
            semester = st.sidebar.text_input("Semester", "VIII")
            academic_year = st.sidebar.text_input("Academic Year", "2025-26")
            year_div = st.sidebar.text_input("Year & Div", "B.TECH")
            assignment_no = st.sidebar.text_input("Assignment No", "03")
            teacher_name = st.sidebar.text_input("Teacher Name")
            max_marks = st.sidebar.text_input("Max Marks", "60")

            subject_incharge = st.sidebar.text_input("Subject Incharge (optional)")
            academic_coordinator = st.sidebar.text_input("Academic Coordinator (optional)")
            hod_name = st.sidebar.text_input("HOD Name (optional)")

            question_count = st.sidebar.slider("Number of Questions", 1, 20, 5)

            st.sidebar.subheader("Bloom Distribution (%)")
            pct_understand = st.sidebar.slider("Understand %", 0, 100, 30)
            pct_apply = st.sidebar.slider("Apply %", 0, 100, 30)
            pct_analyze_eval = st.sidebar.slider("Analyze/Evaluate %", 0, 100, 40)

            total_pct = pct_understand + pct_apply + pct_analyze_eval
            if total_pct == 0:
                pct_understand, pct_apply, pct_analyze_eval = 30, 30, 40
                total_pct = 100

            pct_understand = round((pct_understand / total_pct) * 100)
            pct_apply = round((pct_apply / total_pct) * 100)
            pct_analyze_eval = 100 - pct_understand - pct_apply

            st.sidebar.caption(
                f"Final split: Understand {pct_understand}%, Apply {pct_apply}%, Analyze/Evaluate {pct_analyze_eval}%"
            )

            st.sidebar.subheader("CO / PO / Marks")
            total_cos = st.sidebar.number_input("Total COs", min_value=1, max_value=20, value=6, step=1)
            total_pos = st.sidebar.number_input("Total POs", min_value=1, max_value=20, value=12, step=1)

            m_understand = st.sidebar.number_input("Marks for Understand", min_value=1, max_value=20, value=3, step=1)
            m_apply = st.sidebar.number_input("Marks for Apply", min_value=1, max_value=20, value=5, step=1)
            m_analyze_eval = st.sidebar.number_input("Marks for Analyze/Evaluate", min_value=1, max_value=20, value=7, step=1)

            st.sidebar.subheader("PCU Formatting")
            DOC_FONT_NAME = st.sidebar.selectbox("Font", ["Calibri", "Times New Roman", "Arial"], index=0)
            DOC_FONT_SIZE_PT = st.sidebar.slider("Font Size", 9, 14, 11)
            ROW_HEIGHT_PT = st.sidebar.slider("Row Height (pt)", 16, 40, 24)

            uploaded_file = st.file_uploader("Upload Syllabus (PDF/DOCX)", key="assignment_upload")

            col1, col2 = st.columns(2)

            with col1:
                generate_preview = st.button("👀 Generate Preview", key="assignment_preview")
            with col2:
                clear_preview = st.button("🧹 Clear Preview", key="assignment_clear")

            if clear_preview:
                st.session_state.preview_rows = None
                st.session_state.generated_docx_path = None
                st.rerun()

            if uploaded_file and generate_preview:
                syllabus_text = extract_text(uploaded_file)

                if not subject:
                    st.error("Please enter subject name")
                    st.stop()

                with st.spinner("Generating preview..."):
                    pairs = generate_questions(
                        subject,
                        syllabus_text,
                        question_count,
                        pct_understand,
                        pct_apply,
                        pct_analyze_eval
                    )

                questions_list = [q for (q, b) in pairs][:question_count]
                bloom_labels = [b for (q, b) in pairs][:question_count]

                while len(questions_list) < question_count:
                    questions_list.append("Explain a relevant concept from the given syllabus.")
                    bloom_labels.append("Understand")

                rows = []
                for i, q in enumerate(questions_list, start=1):
                    bloom = bloom_labels[i - 1]
                    rows.append({
                        "Question No.": f"Q{i}",
                        "Question Statement": q,
                        "CO": assign_co(i - 1, int(total_cos)),
                        "PO": assign_po(i - 1, int(total_pos)),
                        "Bloom’s Level": bloom,
                        "Marks": marks_for_bloom(
                            bloom,
                            int(m_understand),
                            int(m_apply),
                            int(m_analyze_eval)
                        )
                    })

                st.session_state.preview_rows = rows

                today = datetime.date.today()
                data = {
                    "{{DEPARTMENT}}": department,
                    "{{SEMESTER}}": semester,
                    "{{ACADEMIC_YEAR}}": academic_year,
                    "{{YEAR_DIV}}": year_div,
                    "{{MAX_MARKS}}": max_marks,
                    "{{SUBJECT}}": subject,
                    "{{ASSIGNMENT_NO}}": assignment_no,
                    "{{TEACHER_NAME}}": teacher_name,
                    "{{DATE}}": today,
                    "{{SUBMISSION_DATE}}": today,
                    "{{SUBJECT_INCHARGE}}": subject_incharge,
                    "{{ACADEMIC_COORDINATOR}}": academic_coordinator,
                    "{{HOD_NAME}}": hod_name
                }

                docx_path = generate_university_docx(
                    data,
                    questions_list,
                    bloom_labels,
                    int(total_cos),
                    int(total_pos),
                    int(m_understand),
                    int(m_apply),
                    int(m_analyze_eval)
                )
                st.session_state.generated_docx_path = docx_path

            if st.session_state.preview_rows:
                st.subheader("📋 Preview (PCU Table)")
                df = pd.DataFrame(st.session_state.preview_rows)
                st.dataframe(df, use_container_width=True)

                st.success("Preview ready. If it looks good, download below 👇")

                if st.session_state.generated_docx_path:
                    with open(st.session_state.generated_docx_path, "rb") as f:
                        st.download_button(
                            label="⬇ Download University Assignment DOCX",
                            data=f,
                            file_name=os.path.basename(st.session_state.generated_docx_path),
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
            else:
                st.info("Upload a syllabus and click **Generate Preview** to see the table before downloading.")

        # =========================
        # QUESTION PAPER MODE
        # =========================
        elif st.session_state.mode == "question_paper":
            st.header("📝 Question Paper Generation")

            st.sidebar.header("Question Paper Details")

            department_qp = st.sidebar.text_input("Department", "CSE", key="qp_department")
            semester_qp = st.sidebar.text_input("Semester", "VIII", key="qp_semester")
            academic_year_qp = st.sidebar.text_input("Academic Year", "2025-26", key="qp_academic_year")

            course_name = st.sidebar.text_input("Course Name", key="qp_course_name")
            course_code = st.sidebar.text_input("Course Code", key="qp_course_code")
            subject_teacher = st.sidebar.text_input("Subject Teacher Name", key="qp_subject_teacher")
            duration = st.sidebar.text_input("Duration", "2 Hours", key="qp_duration")
            total_marks_qp = st.sidebar.text_input("Total Marks", "48", key="qp_total_marks")

            subject_incharge_qp = st.sidebar.text_input("Subject Incharge (optional)", key="qp_subject_incharge")
            academic_coordinator_qp = st.sidebar.text_input("Academic Coordinator (optional)", key="qp_academic_coordinator")
            hod_name_qp = st.sidebar.text_input("HOD Name (optional)", key="qp_hod_name")

            st.sidebar.subheader("CO / PO Mapping")
            total_cos_qp = st.sidebar.number_input("Total COs", min_value=1, max_value=20, value=6, step=1, key="qp_total_cos")
            total_pos_qp = st.sidebar.number_input("Total POs", min_value=1, max_value=20, value=12, step=1, key="qp_total_pos")

            st.subheader("Upload Unit-wise PDFs / DOCX")
            st.caption("Upload one file for each unit")

            unit1_file = st.file_uploader("Upload Unit 1 File", type=["pdf", "docx"], key="unit1")
            unit2_file = st.file_uploader("Upload Unit 2 File", type=["pdf", "docx"], key="unit2")
            unit3_file = st.file_uploader("Upload Unit 3 File", type=["pdf", "docx"], key="unit3")
            unit4_file = st.file_uploader("Upload Unit 4 File", type=["pdf", "docx"], key="unit4")
            unit5_file = st.file_uploader("Upload Unit 5 File", type=["pdf", "docx"], key="unit5")

            st.markdown("### Exam Pattern")
            st.write("**Section A** → Unit 1 → 3 questions printed, solve any 2 → 2 marks each")
            st.write("**Section B** → Unit 2 → 3 questions printed, solve any 2 → 2 marks each")
            st.write("**Section C** → Unit 3 → 3 questions printed, solve any 2 → 5 marks each")
            st.write("**Section D** → Unit 4 → 3 questions printed, solve any 2 → 5 marks each")
            st.write("**Section E** → Unit 5 → 2 questions printed, solve any 1 → 10 marks each")

            col1, col2 = st.columns(2)

            with col1:
                qp_preview = st.button("👀 Generate Question Paper Preview", key="qp_preview")
            with col2:
                qp_clear = st.button("🧹 Clear Question Paper Page", key="qp_clear")

            if qp_clear:
                st.session_state.preview_rows = None
                st.session_state.generated_docx_path = None
                st.rerun()

            if qp_preview:
                missing = []

                if not course_name:
                    missing.append("Course Name")
                if not course_code:
                    missing.append("Course Code")
                if not subject_teacher:
                    missing.append("Subject Teacher Name")
                if not unit1_file:
                    missing.append("Unit 1 File")
                if not unit2_file:
                    missing.append("Unit 2 File")
                if not unit3_file:
                    missing.append("Unit 3 File")
                if not unit4_file:
                    missing.append("Unit 4 File")
                if not unit5_file:
                    missing.append("Unit 5 File")

                if missing:
                    st.error("Please fill/upload: " + ", ".join(missing))
                    st.stop()

                try:
                    with st.spinner("Generating question paper..."):
                        unit1_text = extract_text(unit1_file)
                        unit2_text = extract_text(unit2_file)
                        unit3_text = extract_text(unit3_file)
                        unit4_text = extract_text(unit4_file)
                        unit5_text = extract_text(unit5_file)

                        sec_a = generate_section_questions(
                            course_name, unit1_text, 3, 2, "brief answer", "Understand"
                        )
                        sec_b = generate_section_questions(
                            course_name, unit2_text, 3, 2, "brief answer", "Apply"
                        )
                        sec_c = generate_section_questions(
                            course_name, unit3_text, 3, 5, "descriptive", "Analyze/Evaluate"
                        )
                        sec_d = generate_section_questions(
                            course_name, unit4_text, 3, 5, "descriptive", "Analyze/Evaluate"
                        )
                        sec_e = generate_section_questions(
                            course_name, unit5_text, 2, 10, "long answer", "Analyze/Evaluate"
                        )
                except Exception as e:
                    st.error(f"Question paper generation failed: {str(e)}")
                    st.stop()

                while len(sec_a) < 3:
                    sec_a.append("Explain an important concept from Unit 1.")
                while len(sec_b) < 3:
                    sec_b.append("Explain an important concept from Unit 2.")
                while len(sec_c) < 3:
                    sec_c.append("Discuss an important concept from Unit 3.")
                while len(sec_d) < 3:
                    sec_d.append("Discuss an important concept from Unit 4.")
                while len(sec_e) < 2:
                    sec_e.append("Explain an important concept from Unit 5 in detail.")

                preview_rows = []

                all_sections = [
                    ("A", sec_a, 2, "Understand"),
                    ("B", sec_b, 2, "Apply"),
                    ("C", sec_c, 5, "Analyze/Evaluate"),
                    ("D", sec_d, 5, "Analyze/Evaluate"),
                    ("E", sec_e, 10, "Analyze/Evaluate")
                ]

                q_no = 1
                idx_counter = 0
                for section_name, section_questions, marks, bloom_label in all_sections:
                    for q in section_questions:
                        preview_rows.append({
                            "Section": section_name,
                            "Question No.": f"Q{q_no}",
                            "Question Statement": q,
                            "CO": assign_co(idx_counter, int(total_cos_qp)),
                            "PO": assign_po(idx_counter, int(total_pos_qp)),
                            "Bloom’s Level": bloom_label,
                            "Marks": str(marks)
                        })
                        q_no += 1
                        idx_counter += 1

                st.session_state.preview_rows = preview_rows

                today = datetime.date.today()

                qp_data = {
                    "department": department_qp,
                    "semester": semester_qp,
                    "academic_year": academic_year_qp,
                    "course_name": course_name,
                    "course_code": course_code,
                    "subject_teacher": subject_teacher,
                    "duration": duration,
                    "total_marks": total_marks_qp,
                    "date": today,
                    "exam_date": today,
                    "subject_incharge": subject_incharge_qp,
                    "academic_coordinator": academic_coordinator_qp,
                    "hod_name": hod_name_qp
                }

                qp_docx_path = generate_question_paper_docx(qp_data, preview_rows)
                st.session_state.generated_docx_path = qp_docx_path

            if st.session_state.preview_rows:
                st.subheader("📋 Question Paper Preview")
                df_qp = pd.DataFrame(st.session_state.preview_rows)
                st.dataframe(df_qp, use_container_width=True)

                st.success("Question paper preview ready. Download below 👇")

                if st.session_state.generated_docx_path:
                    with open(st.session_state.generated_docx_path, "rb") as f:
                        st.download_button(
                            label="⬇ Download Question Paper DOCX",
                            data=f,
                            file_name=os.path.basename(st.session_state.generated_docx_path),
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
            else:
                st.info("Upload all 5 unit files and click **Generate Question Paper Preview**.")